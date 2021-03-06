# -*- coding: utf-8 -*-
"""
Created on Thu Apr 05 12:55:11 2018

@author: Alexandre Naaim
"""
import datetime
from param_spt_allfiles import param_spt_allfiles as param_spt_allfiles
from kinematic_allfiles import kinematic_allfiles as kinematic_allfiles
from kinetic_allfiles import kinetic_allfiles as kinetic_allfiles
from tkFileDialog import askopenfilenames, askopenfilename
from tkSimpleDialog import askstring
from tkSimpleDialog import askinteger
import tkMessageBox
from plot_kinematic import plot_kinematic as plot_kinematic
from plot_kinetic import plot_kinetic as plot_kinetic
from plot_kinetic_frontal import plot_kinetic_frontal as plot_kinetic_frontal
from plot_spt import plot_spt as plot_spt
from extract_Schwartz_norm import extract_Schwartz_norm as extract_Schwartz_norm
from plot_emg import plot_emg as plot_emg
import os
import btk
from calculation_extraction_CGM import calculation_extraction_CGM as calculation_extraction_CGM
import json
from extract_SPT_Kid import extract_GaitRite_norm
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from extract_patient import extract_patient as extract_patient
# ------------------------------------------------------------------------------
# Définition des répertoires de travail
# ------------------------------------------------------------------------------
# Repertoire contenant les sujets
data_directory = r'D:\DonneesViconInstallBMF'
# Repertoire ou seront généré les images pour le rapport
report_directory = r'C:\Users\VICON\Desktop\Rapport_Python'

# ------------------------------------------------------------------------------
# Définition de ce que l'utilisateur des choses qu'ils souhaitent faire
# comparaison, cinétique et emg
# ------------------------------------------------------------------------------
test = 1
kinetic_bool = tkMessageBox.askyesno("Title",
                                     "Voulez vous tracer la cinétique?")

emg_bool = tkMessageBox.askyesno("Title",
                                 "Voulez vous tracer les emg ?")
# Initialement c'est deux valeurs sont fausse elles pourront être changer automatiquement
# en cas de détection automatique
emg_case_1_exist = False
emg_case_2_exist = False

old_data = tkMessageBox.askyesno("Title",
                                 "Voulez vous utiliser les données directement calculées dans Nexus ?")
if not old_data:
    there_is_knee_optim = tkMessageBox.askyesno("Title",
                                                "Voulez vous optimiser l'axe du genou pour minimiser le cross_talk?")

else:
    there_is_knee_optim = False
    name_extension_cgm = '_cgm1_1'

detection_autom_bool = tkMessageBox.askyesno("Title",
                                             "Pouvez vous detecter automatiquement les differentes session?")
if not detection_autom_bool:
    comparaison_bool = tkMessageBox.askyesno("Title",
                                             "Voulez vous comparer à d'autres résultats?")

    # Choix des fichiers du premiers repertoire
    filenames_case1 = askopenfilenames(title="Choisir les fichiers de la première condition:",
                                       filetypes=[("Fichiers C3D", "*.c3d")],
                                       initialdir=data_directory)

    # Remise en forme du nom du repertoire pour l'utiliser dans askopenfilename
    one_filename = filenames_case1[0]
    subject_directory_ind_1 = [i for i in range(
        len(one_filename)) if one_filename.startswith('/', i)]
    subject_directory_1 = one_filename[0:subject_directory_ind_1[-1]]
    subject_directory_initial = one_filename[0:subject_directory_ind_1[-1]]

    # choix du fichier statique pour le cas 1
    filenames_stat_case1 = askopenfilename(title="Choisir le fichiers de statique cas 1:",
                                           filetypes=[("Fichiers C3D", "*Cal*.c3d")],
                                           initialdir=subject_directory_1)

    case1_name = askstring("Input", "Quelle est la première condition?")

    if old_data:
        filenames_case1_postCGM = filenames_case1
        subject_directory_1_postCGM = filenames_stat_case1
        extension_pycgm2_case1 = ""
    else:
        [filenames_case1_postCGM, subject_directory_1_postCGM,
         extension_pycgm2_case1] = calculation_extraction_CGM(filenames_stat_case1, filenames_case1, there_is_knee_optim)
        print(extension_pycgm2_case1)

    if comparaison_bool:
        filenames_case2 = askopenfilenames(title="Choisir les fichiers de la deuxième condition:", filetypes=[("Fichiers C3D", "*.c3d")],
                                           initialdir=subject_directory_initial)
        one_filename = filenames_case2[0]

        subject_directory_ind_2 = [i for i in range(
            len(one_filename)) if one_filename.startswith('/', i)]
        subject_directory_2 = one_filename[0:subject_directory_ind_2[-1]]
        # choix du fichier statique pour le cas 1
        filenames_stat_case2 = askopenfilename(title="Choisir le fichiers de statique cas 2:",
                                               filetypes=[("Fichiers C3D", "*Cal*.c3d")],
                                               initialdir=subject_directory_ind_2)
        case2_name = askstring("Input", "Quelle est la deuxième condition?")
        if old_data:
            filenames_case2_postCGM = filenames_case2
            subject_directory_2_postCGM = filenames_stat_case2
            extension_pycgm2_case2 = ""
        else:
            [filenames_case2_postCGM, subject_directory_2_postCGM,
             extension_pycgm2_case2] = calculation_extraction_CGM(filenames_stat_case2, filenames_case2, there_is_knee_optim)
else:

    [case_1_exist, case1_name, filenames_case1, filenames_stat_case1, emg_case_1_exist, emg_filename_case1,
     case_2_exist, case2_name, filenames_case2, filenames_stat_case2, emg_case_2_exist, emg_filename_case2] = extract_patient()

    comparaison_bool = case_1_exist and case_2_exist
    one_filename = filenames_case1[0]
    # repertory from the initial file

    if old_data:
        filenames_case1_postCGM = filenames_case1
        subject_directory_1_postCGM = filenames_stat_case1
        extension_pycgm2_case1 = ""
    else:
        [filenames_case1_postCGM, subject_directory_1_postCGM,
         extension_pycgm2_case1] = calculation_extraction_CGM(filenames_stat_case1, filenames_case1, there_is_knee_optim)
        print(extension_pycgm2_case1)

    if case_2_exist:
        if old_data:
            filenames_case2_postCGM = filenames_case2
            subject_directory_2_postCGM = filenames_stat_case2
            extension_pycgm2_case2 = ""
            name_extension_case2 = '1.1'
        else:
            [filenames_case2_postCGM, subject_directory_2_postCGM,
             extension_pycgm2_case2] = calculation_extraction_CGM(filenames_stat_case2, filenames_case2, there_is_knee_optim)

if not old_data:
    name_extension_cgm = extension_pycgm2_case1
# ------------------------------------------------------------------------------
# Création d'un répertoire ou seront stocké les donnés par sujet
# ------------------------------------------------------------------------------
# Extraction du nom du sujet pour créer un repertoire spécifique ou seront sauvegardé les images
reader = btk.btkAcquisitionFileReader()
reader.SetFilename(str(one_filename))
reader.Update()
acq = reader.GetOutput()

md = acq.GetMetaData()
subject_name = md.FindChild("SUBJECTS").\
    value().FindChild('NAMES').\
    value().GetInfo().ToString()[0]
# Creation of the file containing the data of the patient for today
now = datetime.datetime.now()
date_today = now.strftime("%Y_%m_%d_%Hh%M")
report_directory = os.path.join(report_directory, subject_name.strip(), date_today)

if not os.path.isdir(report_directory):
    os.makedirs(report_directory)

# ------------------------------------------------------------------------------
# Calcul des paramètres spatio temporels et extraction de la cinétique et cinématique
# ------------------------------------------------------------------------------

subject_spt_case1 = param_spt_allfiles(filenames_case1_postCGM)

if kinetic_bool:
    subject_kinematic_case1, subject_kinetic_case1 = kinetic_allfiles(
        filenames_case1_postCGM, extension_pycgm2_case1)
else:
    subject_kinematic_case1 = kinematic_allfiles(filenames_case1_postCGM, extension_pycgm2_case1)

# si l'utilisateurs veut comparer à autre chose on choisit d'autre fichier
if comparaison_bool:
    if extension_pycgm2_case2 != extension_pycgm2_case1:
        tkMessageBox.showerror(
            'Error message', 'Les deux conditions n\'ont pas été traité avec le même modèle. Veuillez retraité les données ou utilisé le bon repertoire!')
        raise ValueError('Les deux conditions n\'ont pas été traité avec le même modèlé')

    # Calcul des paramètres spatio temporels
    subject_spt_case2 = param_spt_allfiles(filenames_case2_postCGM)
    if kinetic_bool:
        subject_kinematic_case2, subject_kinetic_case2 = kinetic_allfiles(
            filenames_case2_postCGM, extension_pycgm2_case2)
    else:
        subject_kinematic_case2 = kinematic_allfiles(
            filenames_case2_postCGM, extension_pycgm2_case2)


subject_kinematic_left_case1 = subject_kinematic_case1["left"]
subject_kinematic_right_case1 = subject_kinematic_case1["right"]

if kinetic_bool:
    subject_kinetic_left_case1 = subject_kinetic_case1["left"]
    subject_kinetic_right_case1 = subject_kinetic_case1["right"]

subject_spt_left_case1 = subject_spt_case1["left"]
subject_spt_right_case1 = subject_spt_case1["right"]

# ------------------------------------------------------------------------------
# Extraction des normes en fonction de la vitesse
# ------------------------------------------------------------------------------
# Calcul de la vitesse pour la norme de Schwartz 2008
bool_vitesse = tkMessageBox.askyesno("Title",
                                     "Voulez vous utiliser une norme adaptée à la vitesse ?")
Age_patient = askinteger("Input", "Quel est l'age du patient ?",
                         minvalue=0, maxvalue=100)
bool_Gender = tkMessageBox.askyesno("Title",
                                    "Est ce une femme ?")
if bool_Gender:
    Gender = 'Female'
else:
    Gender = 'Male'

vitesse_norm = subject_spt_right_case1["mean"]["walking_speed_adm"]
if bool_vitesse:
    if vitesse_norm < 0.227:
        Speed_norm = "VerySlow"
        Speed_norm_REHA = 'C2'
    elif vitesse_norm >= 0.227 and vitesse_norm < 0.363:
        Speed_norm = "Slow"
        Speed_norm_REHA = 'C3'
    elif vitesse_norm >= 0.363 and vitesse_norm < 0.500:
        Speed_norm = "Free"
        Speed_norm_REHA = 'C4'
    elif vitesse_norm >= 0.500 and vitesse_norm < 0.636:
        Speed_norm = "Fast"
        Speed_norm_REHA = 'C5'
    else:
        Speed_norm = "VeryFast"
        Speed_norm_REHA = 'C5'
    bool_vitesse_2 = tkMessageBox.askyesno("Title",
                                           "La norme choisi est :" + Speed_norm +
                                           " .Souhaitez vous la conserver ?")
    if not bool_vitesse_2:
        Speed_norm = "Free"
        Speed_norm_REHA = 'C4'
else:
    Speed_norm = "Free"
    Speed_norm_REHA = 'C4'

[norm_spt, norm_kinematic, norm_kinetic] = extract_Schwartz_norm(Speed=Speed_norm)

with open("SPT_REHA.json", "r") as read_file:
    data = json.load(read_file)
    norm_spt = data[Speed_norm_REHA]

extract_GaitRite_norm(vitesse_norm, Age_patient, Gender)
# ------------------------------------------------------------------------------
# Tracer des graphiques
# ------------------------------------------------------------------------------
# Chaque cote contiendra paramètres spatio temporel, kinematic, kinetics
# extraction des event en list
colorleft_case1 = 'tab:red'
colorright_case1 = 'tab:green'
colorleft_case2 = 'tab:orange'
colorright_case2 = 'tab:blue'

kinematic_pic_S1_LR = plot_kinematic(subject_kinematic_left_case1, subject_spt_left_case1,
                                     colorleft_case1,
                                     subject_kinematic_right_case1, subject_spt_right_case1,
                                     colorright_case1,
                                     norm_spt, norm_kinematic, report_directory,
                                     legend_1="Gauche " + case1_name,
                                     legend_2="Droite " + case1_name,
                                     title="Kinematic_" + case1_name)

SPT_pic_S1_LR, SPT_pic_S1_LR_chiffre = plot_spt(subject_spt_left_case1, colorleft_case1,
                                                subject_spt_right_case1, colorright_case1,
                                                norm_spt, report_directory,
                                                legend_1="Gauche\n" + case1_name,
                                                legend_2="Droite\n" + case1_name,
                                                title="SPT_" + case1_name)

if kinetic_bool:
    kinetic_pic_S1_LR = plot_kinetic(subject_kinetic_left_case1, subject_spt_left_case1,
                                     colorleft_case1,
                                     subject_kinetic_right_case1, subject_spt_right_case1,
                                     colorright_case1,
                                     norm_spt, norm_kinetic, report_directory,
                                     legend_1="Gauche " + case1_name,
                                     legend_2="Droite " + case1_name,
                                     title="Kinetic_" + case1_name)
    kinetic_pic_S1_LR_frontal = plot_kinetic_frontal(subject_kinetic_left_case1, subject_spt_left_case1,
                                                     colorleft_case1,
                                                     subject_kinetic_right_case1, subject_spt_right_case1,
                                                     colorright_case1,
                                                     norm_spt, norm_kinetic, report_directory,
                                                     legend_1="Gauche " + case1_name,
                                                     legend_2="Droite " + case1_name,
                                                     title="Kinetic_frontal" + case1_name)

if emg_bool:
    if not emg_case_1_exist:
        windows_emg_name = "Choisir le fichies de tracer EMG condition " + \
            case1_name + " :"
        emg_filename_case1 = askopenfilenames(title=windows_emg_name,
                                              filetypes=[("Fichiers C3D", "*.c3d")],
                                              initialdir=subject_directory_1_postCGM)
        emg_plot_case1 = plot_emg(emg_filename_case1[0], colorleft_case1, colorright_case1,
                                  report_directory, title="EMG " + case1_name)
    else:
        # On sait ici que emg_filename_case1 est une chaine de caractèere et pas une liste
        emg_plot_case1 = plot_emg(emg_filename_case1, colorleft_case1, colorright_case1,
                                  report_directory, title="EMG " + case1_name)


if comparaison_bool:
    if emg_bool:
        if not emg_case_2_exist:
            windows_emg_name = "Choisir le fichies de tracer EMG condition " + \
                case2_name + " :"
            emg_filename_case2 = askopenfilenames(title=windows_emg_name,
                                                  filetypes=[("Fichiers C3D", "*.c3d")],
                                                  initialdir=subject_directory_2_postCGM)
            emg_plot_case2 = plot_emg(emg_filename_case2[0], colorleft_case2, colorright_case2,
                                      report_directory, title="EMG " + case2_name)
        else:
            # On sait ici que emg_filename_case1 est une chaine de caractèere et pas une liste
            emg_plot_case2 = plot_emg(emg_filename_case2, colorleft_case2, colorright_case2,
                                      report_directory, title="EMG " + case2_name)

    subject_kinematic_left_case2 = subject_kinematic_case2["left"]
    subject_kinematic_right_case2 = subject_kinematic_case2["right"]
    if kinetic_bool:
        subject_kinetic_left_case2 = subject_kinetic_case2["left"]
        subject_kinetic_right_case2 = subject_kinetic_case2["right"]

    subject_spt_left_case2 = subject_spt_case2["left"]
    subject_spt_right_case2 = subject_spt_case2["right"]

    kinematic_pic_S2_LR = plot_kinematic(subject_kinematic_left_case2, subject_spt_left_case2,
                                         colorleft_case2,
                                         subject_kinematic_right_case2, subject_spt_right_case2,
                                         colorright_case2,
                                         norm_spt, norm_kinematic, report_directory,
                                         legend_1="Gauche " + case2_name,
                                         legend_2="Droite " + case2_name,
                                         title="Kinematic_" + case2_name)

    SPT_pic_S2_LR, SPT_pic_S2_LR_chiffre = plot_spt(subject_spt_left_case2, colorleft_case2,
                                                    subject_spt_right_case2, colorright_case2,
                                                    norm_spt, report_directory,
                                                    legend_1="Gauche\n" + case2_name,
                                                    legend_2="Droite\n" + case2_name,
                                                    title="SPT_" + case2_name)

    kinematic_pic_Left = plot_kinematic(subject_kinematic_left_case1, subject_spt_left_case1,
                                        colorleft_case1,
                                        subject_kinematic_left_case2, subject_spt_left_case2,
                                        colorleft_case2,
                                        norm_spt, norm_kinematic, report_directory,
                                        legend_1="Gauche " + case1_name,
                                        legend_2="Gauche " + case2_name,
                                        title="Kinematic_Comparaison_Left")

    SPT_pic_Left, SPT_pic_Left_chiffre = plot_spt(subject_spt_left_case1, colorleft_case1,
                                                  subject_spt_left_case2, colorleft_case2,
                                                  norm_spt, report_directory,
                                                  legend_1="Gauche\n" + case1_name,
                                                  legend_2="Gauche\n" + case2_name,
                                                  title="SPT_Comparaison_Left")

    kinematic_pic_Right = plot_kinematic(subject_kinematic_right_case1, subject_spt_right_case1,
                                         colorright_case1,
                                         subject_kinematic_right_case2, subject_spt_right_case2,
                                         colorright_case2,
                                         norm_spt, norm_kinematic, report_directory,
                                         legend_1="Droite " + case1_name,
                                         legend_2="Droite " + case2_name,
                                         title="Kinematic_Comparaison_Right")

    SPT_pic_Right, SPT_pic_Right_chiffre = plot_spt(subject_spt_right_case1, colorright_case1,
                                                    subject_spt_right_case2, colorright_case2,
                                                    norm_spt, report_directory,
                                                    legend_1="Droite\n" + case1_name,
                                                    legend_2="Droite\n" + case2_name,
                                                    title="SPT_Comparaison_Right")

    if kinetic_bool:
        kinetic_pic_S2_LR = plot_kinetic(subject_kinetic_left_case2, subject_spt_left_case2,
                                         colorleft_case2,
                                         subject_kinetic_right_case2, subject_spt_right_case2,
                                         colorright_case2,
                                         norm_spt, norm_kinetic, report_directory,
                                         legend_1="Gauche " + case2_name,
                                         legend_2="Droite " + case2_name,
                                         title="Kinetic_" + case2_name)
        kinetic_pic_S2_LR_frontal = plot_kinetic_frontal(subject_kinetic_left_case2, subject_spt_left_case2,
                                                         colorleft_case2,
                                                         subject_kinetic_right_case2, subject_spt_right_case2,
                                                         colorright_case2,
                                                         norm_spt, norm_kinetic, report_directory,
                                                         legend_1="Gauche " + case2_name,
                                                         legend_2="Droite " + case2_name,
                                                         title="Kinetic_Frontal" + case2_name)

        kinetic_pic_Left = plot_kinetic(subject_kinetic_left_case1, subject_spt_left_case1,
                                        colorleft_case1,
                                        subject_kinetic_left_case2, subject_spt_left_case2,
                                        colorleft_case2,
                                        norm_spt, norm_kinetic, report_directory,
                                        legend_1="Gauche " + case1_name,
                                        legend_2="Gauche " + case2_name,
                                        title="Kinetic_Comparaison_Left")
        kinetic_pic_Left_frontal = plot_kinetic_frontal(subject_kinetic_left_case1, subject_spt_left_case1,
                                                        colorleft_case1,
                                                        subject_kinetic_left_case2, subject_spt_left_case2,
                                                        colorleft_case2,
                                                        norm_spt, norm_kinetic, report_directory,
                                                        legend_1="Gauche " + case1_name,
                                                        legend_2="Gauche " + case2_name,
                                                        title="Kinetic_Comparaison_Left_Frontal")

        kinetic_pic_Right = plot_kinetic(subject_kinetic_right_case1, subject_spt_right_case1,
                                         colorright_case1,
                                         subject_kinetic_right_case2, subject_spt_right_case2,
                                         colorright_case2,
                                         norm_spt, norm_kinetic, report_directory,
                                         legend_1="Droite " + case1_name,
                                         legend_2="Droite " + case2_name,
                                         title="Kinetic_Comparaison_Right")
        kinetic_pic_Right_frontal = plot_kinetic_frontal(subject_kinetic_right_case1, subject_spt_right_case1,
                                                         colorright_case1,
                                                         subject_kinetic_right_case2, subject_spt_right_case2,
                                                         colorright_case2,
                                                         norm_spt, norm_kinetic, report_directory,
                                                         legend_1="Droite " + case1_name,
                                                         legend_2="Droite " + case2_name,
                                                         title="Kinetic_Comparaison_Right_Frontal")
# Generation du rapport word


def newpage_report(name_param, pic_1):
    document.add_page_break()
    document.add_heading(name_param, level=1)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(10)
    run = paragraph.add_run()
    run.add_picture(pic_1, width=Inches(5))
    document.add_heading(unicode('Interprétation :', 'utf-8'), level=3)


def newpage_report_2images(name_param, pic_1, pic_2):
    document.add_page_break()
    document.add_heading(name_param, level=1)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(10)
    run = paragraph.add_run()
    run.add_picture(pic_1, width=Inches(4))
    run.add_picture(pic_2, width=Inches(4))
    document.add_heading(unicode('Interprétation :', 'utf-8'), level=3)


def newpage_report_comp(name_param, pic_1, pic_2):
    document.add_page_break()
    document.add_heading(name_param, level=1)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(10)
    run = paragraph.add_run()
    run.add_picture(pic_1, width=Inches(3.4))
    run.add_picture(pic_2, width=Inches(3.4))
    document.add_heading(unicode('Interprétation :', 'utf-8'), level=3)


def newpage_report_comp_4images(name_param, pic_1, pic_2, pic_3, pic_4):
    document.add_page_break()
    document.add_heading(name_param, level=1)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(10)
    run = paragraph.add_run()
    run.add_picture(pic_1, width=Inches(3.2))
    run.add_picture(pic_2, width=Inches(3.2))
    run.add_picture(pic_3, width=Inches(3.2))
    run.add_picture(pic_4, width=Inches(3.2))
    document.add_heading(unicode('Interprétation :', 'utf-8'), level=3)


document = Document('Rapport_massue_Vierge.docx')

if comparaison_bool:
    p = document.add_paragraph(unicode('Le conventional gait model ', 'utf-8'))
    name_final = name_extension_cgm.replace('_', '.')
    p.add_run(unicode(name_final[4:], 'utf-8'))
    p.add_run(unicode(' a été utilisé (', 'utf-8'))
    paragraph_format = p.paragraph_format
    p.add_run('https://pycgm2.github.io/').italic = True
    p.add_run(').')

    if there_is_knee_optim:
        p2 = document.add_paragraph(
            unicode('La méthode de réorientation du genou DynaKad a été appliquée (', 'utf-8'))
        paragraph_format = p2.paragraph_format
        paragraph_format.space_before = Pt(10)
        document.add_paragraph('\n')
        p2.add_run(unicode(
            'Baker, R., 1999. A new approach to determine the hip rotation from clinical gait analysis data 18, 655–667', 'utf-8')).italic = True
        document.add_run(unicode(').', 'utf-8'))

    p3 = document.add_paragraph()
    p3.add_run('Objectifs de la demande :').bold = True

    newpage_report_comp_4images(unicode('Paramètres Spatio-temporels', 'utf-8'),
                                SPT_pic_S1_LR_chiffre, SPT_pic_S2_LR_chiffre, SPT_pic_S1_LR, SPT_pic_S2_LR)
    newpage_report_comp_4images(unicode('Paramètres Spatio-temporels',
                                        'utf-8'), SPT_pic_Left_chiffre, SPT_pic_Right_chiffre, SPT_pic_Left, SPT_pic_Right)

    newpage_report_comp(unicode('Cinématique', 'utf-8'), kinematic_pic_S1_LR, kinematic_pic_S2_LR)
    if kinetic_bool:
        #newpage_report_comp(unicode('Cinétique :', 'utf-8'), kinetic_pic_S1_LR, kinetic_pic_S2_LR)
        newpage_report_comp_4images(unicode('Cinétique :', 'utf-8'), kinetic_pic_S1_LR,
                                    kinetic_pic_S2_LR, kinetic_pic_S1_LR_frontal, kinetic_pic_S2_LR_frontal)
    newpage_report_comp(unicode('Cinématique coté par coté : ', 'utf-8'),
                        kinematic_pic_Left, kinematic_pic_Right)
    if kinetic_bool:
        newpage_report_comp_4images(unicode('Cinétique coté par coté : ', 'utf-8'),
                                    kinetic_pic_Left, kinetic_pic_Right, kinetic_pic_Left_frontal, kinetic_pic_Right_frontal)

    if emg_bool:
        newpage_report_comp(unicode('EMG', 'utf-8'), emg_plot_case1, emg_plot_case2)

else:
    newpage_report_2images(unicode('Paramètres Spatio-temporels', 'utf-8'),
                           SPT_pic_S1_LR_chiffre, SPT_pic_S1_LR)
    newpage_report(unicode('Cinématique', 'utf-8'), kinematic_pic_S1_LR)
    if kinetic_bool:
        newpage_report_2images(unicode('Cinétique', 'utf-8'),
                               kinetic_pic_S1_LR, kinetic_pic_S1_LR_frontal)

    if emg_bool:
        newpage_report(unicode('EMG', 'utf-8'), emg_plot_case1)


document.save(os.path.join(report_directory, 'rapport.docx'))
